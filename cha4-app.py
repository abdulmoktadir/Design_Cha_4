import streamlit as st
import numpy as np
import pandas as pd
import graphviz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64
import matplotlib.pyplot as plt
import re
import hmac
from pathlib import Path

# ============================================================
# PAGE CONFIG
# ============================================================
st.set_page_config(
    page_title="IT2TrFS MCDM Toolkit | Delphi · WINGS · CoCoSo",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ============================================================
# CSS STYLING
# ============================================================
st.markdown("""
<style>
    /* Import modern font */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

    * {
        font-family: 'Inter', sans-serif;
    }

    /* Main background */
    .stApp {
        background-color: #f8fafc;
    }

    /* ===== Main workspace banner (right side / main body) ===== */
    .workspace-banner {
        background: linear-gradient(135deg, #0f172a 0%, #1d4ed8 50%, #06b6d4 100%);
        border-radius: 22px;
        padding: 1.35rem 1.45rem;
        margin-bottom: 1.1rem;
        box-shadow: 0 14px 32px rgba(29, 78, 216, 0.18);
        border: 1px solid rgba(255,255,255,0.18);
    }
    .workspace-banner-title {
        color: #ffffff;
        font-size: 1.75rem;
        font-weight: 800;
        margin: 0 0 0.35rem 0;
        letter-spacing: -0.03em;
    }
    .workspace-banner-subtitle {
        color: rgba(255,255,255,0.92);
        font-size: 0.98rem;
        line-height: 1.55;
        margin: 0 0 0.95rem 0;
    }
    .workspace-chip-wrap {
        display: flex;
        gap: 0.55rem;
        flex-wrap: wrap;
    }
    .workspace-chip {
        display: inline-block;
        font-size: 0.78rem;
        font-weight: 700;
        padding: 0.34rem 0.78rem;
        border-radius: 999px;
        background: rgba(255,255,255,0.14);
        color: #ffffff;
        border: 1px solid rgba(255,255,255,0.20);
        backdrop-filter: blur(5px);
    }

    /* ===== Sidebar expander styling: blue / grey ===== */
    [data-testid="stSidebar"] div[data-testid="stExpander"] {
        border: 1px solid rgba(148, 163, 184, 0.22);
        border-radius: 14px;
        background: rgba(148, 163, 184, 0.08);
        overflow: hidden;
        margin-bottom: 0.75rem;
    }
    [data-testid="stSidebar"] div[data-testid="stExpander"] details {
        border-radius: 14px;
    }

    [data-testid="stSidebar"] div[data-testid="stExpander"] summary {
        background: linear-gradient(135deg, rgba(59,130,246,0.22), rgba(71,85,105,0.34));
        border-radius: 12px;
        padding: 0.2rem 0.35rem;
    }

    [data-testid="stSidebar"] div[data-testid="stExpander"] summary:hover {
        background: linear-gradient(135deg, rgba(59,130,246,0.30), rgba(100,116,139,0.40));
    }

    [data-testid="stSidebar"] div[data-testid="stExpander"] summary p,
    [data-testid="stSidebar"] div[data-testid="stExpander"] summary span {
        color: #e2e8f0 !important;
        font-weight: 700;
    }

    [data-testid="stSidebar"] div[data-testid="stExpander"] details[open] > summary {
        background: linear-gradient(135deg, rgba(37,99,235,0.30), rgba(71,85,105,0.45));
    }

    /* Sidebar */
    [data-testid="stSidebar"] {
        background-color: #0f172a;
        border-right: none;
    }
    [data-testid="stSidebar"] .sidebar-content {
        padding: 1.5rem 1rem;
    }
    [data-testid="stSidebar"] h1,
    [data-testid="stSidebar"] h2,
    [data-testid="stSidebar"] h3,
    [data-testid="stSidebar"] .stMarkdown,
    [data-testid="stSidebar"] .stRadio label,
    [data-testid="stSidebar"] .stButton > button {
        color: #f1f5f9 !important;
    }
    [data-testid="stSidebar"] .stRadio label div {
        color: #f1f5f9;
    }
    [data-testid="stSidebar"] .stRadio label:hover {
        background-color: #1e293b;
        border-radius: 8px;
    }
    [data-testid="stSidebar"] .stSuccess {
        background-color: #1e293b;
        color: #e2e8f0;
    }
    [data-testid="stSidebar"] .stButton > button {
        background-color: transparent;
        border: 1px solid #94a3b8;
        color: #f1f5f9;
        border-radius: 8px;
    }
    [data-testid="stSidebar"] .stButton > button:hover {
        background-color: #3b82f6;
        border-color: #3b82f6;
    }

    /* Cards – subtle and clean */
    .app-card, .inventor-card {
        background: #ffffff;
        border-radius: 16px;
        padding: 1.25rem 1.5rem;
        margin-bottom: 1.5rem;
        border: 1px solid #eef2f6;
        transition: all 0.2s;
    }
    .app-card h4, .inventor-card h4 {
        margin-top: 0;
        margin-bottom: 0.75rem;
        color: #1e2a3e;
        font-weight: 600;
        font-size: 1.1rem;
    }

    /* Researcher profiles */
    .inventor-name {
        font-size: 1.3rem;
        font-weight: 700;
        color: #0f172a;
        margin: 0.5rem 0 0.25rem;
    }
    .inventor-role {
        font-size: 0.9rem;
        color: #475569;
        margin-bottom: 0.5rem;
    }
    .inventor-mini {
        font-size: 0.85rem;
        color: #334155;
        margin-top: 0.5rem;
    }
    .inventor-divider {
        height: 2px;
        background: linear-gradient(90deg, #e2e8f0, #94a3b8, #e2e8f0);
        margin: 2rem 0 1rem;
    }
    .inventor-section-title {
        font-size: 1.6rem;
        font-weight: 600;
        margin-bottom: 1.2rem;
        color: #0f172a;
        text-align: center;
    }

    /* Headers */
    h1, h2, h3 {
        color: #0f172a;
        font-weight: 600;
        letter-spacing: -0.3px;
    }
    h1 { font-size: 2rem; margin-bottom: 0.25rem; }
    h2 { font-size: 1.5rem; margin-top: 1.5rem; margin-bottom: 1rem; }
    h3 { font-size: 1.25rem; }

    /* Buttons */
    .stButton > button {
        border-radius: 8px;
        font-weight: 500;
        padding: 0.5rem 1rem;
        transition: all 0.2s;
    }
    .stButton > button[kind="primary"] {
        background-color: #3b82f6;
        border: none;
        color: white;
    }
    .stButton > button[kind="primary"]:hover {
        background-color: #2563eb;
        box-shadow: 0 2px 6px rgba(59,130,246,0.2);
    }

    /* DataFrames & Editors */
    .stDataFrame, .stDataEditor {
        border-radius: 12px;
        overflow: hidden;
        border: 1px solid #e2e8f0;
    }
    .stDataFrame thead tr th, .stDataEditor thead tr th {
        background-color: #f8fafc;
        font-weight: 600;
        color: #1e2a3e;
        padding: 0.75rem 1rem;
    }

    /* Metrics */
    .stMetric {
        background: #f8fafc;
        border-radius: 12px;
        padding: 0.8rem;
        text-align: center;
        border: 1px solid #eef2f6;
    }

    /* Footer */
    .app-footer {
        text-align: center;
        margin-top: 2.5rem;
        padding: 1rem;
        color: #64748b;
        font-size: 0.75rem;
        border-top: 1px solid #e2e8f0;
    }

        /* Login page - compact and centered */
    .login-page {
        max-width: 520px;
        margin: 2.5rem auto 0 auto;
        padding: 0 0.5rem 1rem 0.5rem;
    }

    .login-hero {
        text-align: center;
        margin-bottom: 1.25rem;
    }

    .login-hero h1 {
        font-size: 2.1rem;
        margin: 0 0 0.45rem 0;
        color: #0f172a;
        letter-spacing: -0.03em;
    }

    .login-hero p {
        margin: 0;
        color: #475569;
        font-size: 0.98rem;
        line-height: 1.55;
    }

    .login-form-note {
        text-align: center;
        color: #64748b;
        font-size: 0.86rem;
        margin-bottom: 0.9rem;
    }

    /* Style the Streamlit form itself as the card */
    div[data-testid="stForm"] {
        background: #ffffff;
        border: 1px solid #e2e8f0;
        border-radius: 18px;
        padding: 1.2rem 1.2rem 0.9rem 1.2rem;
        box-shadow: 0 12px 28px rgba(15, 23, 42, 0.06);
    }

    div[data-testid="stForm"] [data-testid="stFormSubmitButton"] > button {
        width: 100%;
        margin-top: 0.35rem;
    }

    /* ===== Premium sidebar header ===== */
    .sidebar-brand {
        background: linear-gradient(135deg, #111827, #1e293b);
        border: 1px solid rgba(148, 163, 184, 0.18);
        border-radius: 18px;
        padding: 1rem 1rem 0.9rem 1rem;
        margin-bottom: 1rem;
        box-shadow: 0 10px 30px rgba(0,0,0,0.18);
    }
    .sidebar-brand h2 {
        color: #f8fafc !important;
        font-size: 1.1rem;
        margin: 0 0 0.25rem 0;
        letter-spacing: -0.2px;
    }
    .sidebar-brand p {
        color: #cbd5e1 !important;
        font-size: 0.82rem;
        line-height: 1.45;
        margin: 0;
    }
    .sidebar-chip-wrap {
        display: flex;
        gap: 0.4rem;
        flex-wrap: wrap;
        margin-top: 0.75rem;
    }
    .sidebar-chip {
        display: inline-block;
        font-size: 0.72rem;
        font-weight: 600;
        padding: 0.24rem 0.6rem;
        border-radius: 999px;
        background: rgba(59, 130, 246, 0.16);
        color: #dbeafe !important;
        border: 1px solid rgba(96, 165, 250, 0.25);
    }

    /* ===== Sidebar profile section ===== */
    .sidebar-section-title {
        color: #e2e8f0 !important;
        font-size: 0.88rem;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 0.08em;
        margin: 0.25rem 0 0.75rem 0;
    }
    .sidebar-profile-card {
        background: rgba(255, 255, 255, 0.06);
        border: 1px solid rgba(148, 163, 184, 0.14);
        border-radius: 16px;
        padding: 0.85rem;
        margin-bottom: 0.85rem;
        backdrop-filter: blur(6px);
    }
    .sidebar-profile-name {
        color: #f8fafc;
        font-size: 0.95rem;
        font-weight: 700;
        margin-top: 0.55rem;
        margin-bottom: 0.2rem;
        line-height: 1.35;
    }
    .sidebar-profile-role {
        color: #cbd5e1;
        font-size: 0.78rem;
        line-height: 1.35;
        margin-bottom: 0.45rem;
    }
    .sidebar-profile-text {
        color: #94a3b8;
        font-size: 0.76rem;
        line-height: 1.5;
        margin-bottom: 0.45rem;
    }

    /* ===== Better navigation polish ===== */
    [data-testid="stSidebar"] .stRadio > div {
        gap: 0.25rem;
    }
    [data-testid="stSidebar"] .stRadio label {
        padding: 0.45rem 0.55rem;
        border-radius: 10px;
        transition: all 0.2s ease;
    }

    /* ===== Softer main cards ===== */
    .app-card {
        box-shadow: 0 8px 24px rgba(15, 23, 42, 0.04);
    }
    .app-card:hover, .inventor-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 10px 26px rgba(15, 23, 42, 0.08);
    }
    /* ===== Colorful page title box ===== */
    .page-banner {
        background: linear-gradient(135deg, #2563eb 0%, #7c3aed 55%, #ec4899 100%);
        border-radius: 20px;
        padding: 1.2rem 1.4rem;
        margin-bottom: 1.15rem;
        box-shadow: 0 12px 30px rgba(59, 130, 246, 0.18);
        border: 1px solid rgba(255,255,255,0.18);
    }
    .page-banner-title {
        color: white;
        font-size: 1.65rem;
        font-weight: 800;
        margin: 0 0 0.3rem 0;
        letter-spacing: -0.02em;
    }
    .page-banner-subtitle {
        color: rgba(255,255,255,0.92);
        font-size: 0.96rem;
        line-height: 1.5;
        margin: 0;
    }
    
    /* Optional alternate banner themes */
    .page-banner.green {
        background: linear-gradient(135deg, #059669 0%, #0ea5e9 55%, #2563eb 100%);
    }
    .page-banner.orange {
        background: linear-gradient(135deg, #f59e0b 0%, #ef4444 55%, #ec4899 100%);
    }
    .page-banner.slate {
        background: linear-gradient(135deg, #0f172a 0%, #1d4ed8 55%, #06b6d4 100%);
    }
</style>
""", unsafe_allow_html=True)

# ============================================================
# AUTHENTICATION
# ============================================================
def logout():
    st.session_state.authenticated = False
    st.rerun()


def check_password():
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False

    if st.session_state.authenticated:
        return True

    expected_password = st.secrets.get("APP_PASSWORD", None)

    left, center, right = st.columns([1.15, 1.3, 1.15])

    with center:
        st.markdown(
            """
            <div class="login-page">
                <div class="login-hero">
                    <h1>IT2TrFS MCDM Toolkit</h1>
                    <p>
                        Secure access to a decision analytics workspace for
                        IT2TrFS Delphi, WINGS &amp; CoCoSo.
                    </p>
                </div>
                <div class="login-form-note">
                    Sign in with the application password
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        if expected_password is None:
            st.error(
                "APP_PASSWORD is not configured. Add it in Streamlit secrets "
                "(.streamlit/secrets.toml locally or Settings → Secrets on Streamlit Cloud)."
            )
            return False

        with st.form("login_form", clear_on_submit=False):
            password = st.text_input(
                "Password",
                type="password",
                placeholder="Enter application password",
            )
            submitted = st.form_submit_button("Log in", use_container_width=True)

        if submitted:
            if hmac.compare_digest(password, str(expected_password)):
                st.session_state.authenticated = True
                st.success("Access granted.")
                st.rerun()
            else:
                st.error("Incorrect password.")

    return False


# ============================================================
# RESEARCHER PROFILE HELPERS
# ============================================================
def get_asset_path(filename):
    """Resolve image file path from 'assets' folder."""
    return Path(__file__).parent / "assets" / filename


def render_sidebar_profile_card(name, role, institution, image_path, brief_text, full_bio=None, extras=None):
    st.markdown('<div class="sidebar-profile-card">', unsafe_allow_html=True)

    if Path(image_path).exists():
        st.image(str(image_path), use_container_width=True)
    else:
        st.caption(f"Image not found: {image_path}")

    st.markdown(
        f"""
        <div class="sidebar-profile-name">{name}</div>
        <div class="sidebar-profile-role">{role}<br>{institution}</div>
        <div class="sidebar-profile-text">{brief_text}</div>
        """,
        unsafe_allow_html=True,
    )

    if extras:
        for item in extras:
            st.markdown(
                f'<div class="sidebar-profile-text">• {item}</div>',
                unsafe_allow_html=True,
            )

    if full_bio:
        with st.expander(f"More about {name}", expanded=False):
            st.write(full_bio)

    st.markdown("</div>", unsafe_allow_html=True)


def render_sidebar_research_profiles():
    st.sidebar.markdown("---")
    st.sidebar.markdown(
        '<div class="sidebar-section-title">Research Profiles</div>',
        unsafe_allow_html=True,
    )

    with st.sidebar.expander("👥 View researcher profiles", expanded=False):
        render_sidebar_profile_card(
            name="Prof. J.Z. Ren 任競爭",
            role="Associate Professor",
            institution="The Hong Kong Polytechnic University",
            image_path=get_asset_path("prof_jz_ren.png"),
            brief_text=(
                "Process systems engineering for energy, environment, "
                "and sustainability. Recipient of the 2022 APEC ASPIRE Prize."
            ),
            full_bio=(
                "Dr. Jingzheng Ren is currently an Associate Professor at The Hong Kong "
                "Polytechnic University. His research focuses on process systems engineering "
                "for energy, environment and sustainability, including innovative industrial "
                "processes, decision-making tools, and optimization models for sustainable "
                "and carbon-neutral industrial systems. He has published extensively and has "
                "received major international recognition, including the 2022 APEC Science "
                "Prize for Innovation, Research and Education (ASPIRE Prize)."
            ),
        )

        render_sidebar_profile_card(
            name="Md. Abdul Moktadir",
            role="Assistant Professor (Leather Products Engineering)",
            institution="University of Dhaka / PolyU Presidential PhD Fellow",
            image_path=get_asset_path("abdul_moktadir.png"),
            brief_text=(
                "Research interests include sustainable supply chains, logistics, "
                "risk management, Industry 4.0, and circular economy."
            ),
            extras=[
                "Affiliation: University of Dhaka",
                "Program: PolyU Presidential PhD Fellow",
            ],
            full_bio=(
                "Md. Abdul Moktadir is currently pursuing a PhD in Industrial and Systems "
                "Engineering at The Hong Kong Polytechnic University. He also serves as an "
                "Assistant Professor of Leather Products Engineering at the University of "
                "Dhaka. His work has appeared in several international journals, and his "
                "research interests include sustainable supply chain management, risk "
                "management, energy-efficient supply chain planning and design, logistics, "
                "Industry 4.0, and circular economy."
            ),
        )

def render_footer():
    st.markdown(
        """
        <div class="app-footer">
            © 2026 Developed by <strong>Moktadir M.A.</strong> and <strong>REN J.Z.</strong>
        </div>
        """,
        unsafe_allow_html=True,
    )
def render_page_banner(title, subtitle, theme="default"):
    theme_class = "" if theme == "default" else f" {theme}"
    st.markdown(
        f"""
        <div class="page-banner{theme_class}">
            <div class="page-banner-title">{title}</div>
            <div class="page-banner-subtitle">{subtitle}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
def render_workspace_banner():
    st.markdown(
        """
        <div class="workspace-banner">
            <div class="workspace-banner-title">🧠 IT2TrFS MCDM Suite</div>
            <div class="workspace-banner-subtitle">
                A professional decision analytics workspace for Delphi, WINGS, and CoCoSo.
            </div>
            <div class="workspace-chip-wrap">
                <span class="workspace-chip">Delphi</span>
                <span class="workspace-chip">WINGS</span>
                <span class="workspace-chip">CoCoSo</span>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

def dataframe_dict_to_excel_bytes(sheet_map):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        for sheet_name, df in sheet_map.items():
            safe_name = re.sub(r'[\\/*?:\\[\\]]', "", str(sheet_name))[:31] or "Sheet1"
            export_df = df.copy()
            export_df.to_excel(writer, index=False, sheet_name=safe_name)
    output.seek(0)
    return output.getvalue()

# ============================================================
# IT2TrFS REPRESENTATION
# ============================================================
def format_it2(it2):
    u, l = it2
    return (
        f"(({u[0]:.6f},{u[1]:.6f},{u[2]:.6f},{u[3]:.6f};{u[4]:.1f},{u[5]:.1f}), "
        f"({l[0]:.6f},{l[1]:.6f},{l[2]:.6f},{l[3]:.6f};{l[4]:.1f},{l[5]:.1f}))"
    )

def zero_it2():
    return ((0, 0, 0, 0, 1, 1), (0, 0, 0, 0, 0.9, 0.9))

def add_it2(A, B):
    Au, Al = A
    Bu, Bl = B
    new_u = (
        Au[0] + Bu[0], Au[1] + Bu[1], Au[2] + Bu[2], Au[3] + Bu[3],
        min(Au[4], Bu[4]), min(Au[5], Bu[5])
    )
    new_l = (
        Al[0] + Bl[0], Al[1] + Bl[1], Al[2] + Bl[2], Al[3] + Bl[3],
        min(Al[4], Bl[4]), min(Al[5], Bl[5])
    )
    return (new_u, new_l)

def sub_it2(A, B):
    Au, Al = A
    Bu, Bl = B
    new_u = (
        Au[0] - Bu[0], Au[1] - Bu[1], Au[2] - Bu[2], Au[3] - Bu[3],
        min(Au[4], Bu[4]), min(Au[5], Bu[5])
    )
    new_l = (
        Al[0] - Bl[0], Al[1] - Bl[1], Al[2] - Bl[2], Al[3] - Bl[3],
        min(Al[4], Bl[4]), min(Al[5], Bl[5])
    )
    return (new_u, new_l)

def mul_it2(A, B):
    Au, Al = A
    Bu, Bl = B
    new_u = (
        Au[0] * Bu[0], Au[1] * Bu[1], Au[2] * Bu[2], Au[3] * Bu[3],
        min(Au[4], Bu[4]), min(Au[5], Bu[5])
    )
    new_l = (
        Al[0] * Bl[0], Al[1] * Bl[1], Al[2] * Bl[2], Al[3] * Bl[3],
        min(Al[4], Bl[4]), min(Al[5], Bl[5])
    )
    return (new_u, new_l)

def scalar_mul_it2(k, A):
    Au, Al = A
    new_u = (k * Au[0], k * Au[1], k * Au[2], k * Au[3], Au[4], Au[5])
    new_l = (k * Al[0], k * Al[1], k * Al[2], k * Al[3], Al[4], Al[5])
    return (new_u, new_l)

def it2_pow(A, w):
    Au, Al = A
    def pw(x):
        if x == 0 and w == 0:
            return 1.0
        return float(x) ** float(w)
    new_u = (pw(Au[0]), pw(Au[1]), pw(Au[2]), pw(Au[3]), Au[4], Au[5])
    new_l = (pw(Al[0]), pw(Al[1]), pw(Al[2]), pw(Al[3]), Al[4], Al[5])
    return (new_u, new_l)

def cocoso_crisp_score(it2):
    Au, Al = it2
    a, b, c, d, uh1, uh2 = Au
    e, f, g, h, lh1, lh2 = Al
    score_u = (((d - a) + ((uh2 * c) - a) + ((uh1 * b) - a)) / 4.0) + a
    score_l = (((h - e) + ((lh2 * g) - e) + ((lh1 * f) - e)) / 4.0) + e
    return (score_u + score_l) / 2.0

def it2_score_components(it2):
    Au, Al = it2
    a, b, c, d, uh1, uh2 = Au
    e, f, g, h, lh1, lh2 = Al
    score_u = (((d - a) + ((uh2 * c) - a) + ((uh1 * b) - a)) / 4.0) + a
    score_l = (((h - e) + ((lh2 * g) - e) + ((lh1 * f) - e)) / 4.0) + e
    crisp = (score_u + score_l) / 2.0
    return score_u, score_l, crisp

# ============================================================
# IT2TrFS-CoCoSo linguistic scale & helpers
# ============================================================
COCOSO_LINGUISTIC_TERMS = {
    "VP": ((0, 0, 0, 0.1, 1, 1), (0.05, 0, 0, 0.05, 0.9, 0.9)),
    "P": ((0, 0.1, 0.1, 0.3, 1, 1), (0.05, 0.1, 0.1, 0.25, 0.9, 0.9)),
    "MP": ((0.1, 0.3, 0.3, 0.5, 1, 1), (0.15, 0.3, 0.3, 0.45, 0.9, 0.9)),
    "F": ((0.3, 0.5, 0.5, 0.7, 1, 1), (0.35, 0.5, 0.5, 0.65, 0.9, 0.9)),
    "MG": ((0.5, 0.7, 0.7, 0.9, 1, 1), (0.55, 0.7, 0.7, 0.85, 0.9, 0.9)),
    "G": ((0.7, 0.9, 0.9, 1.0, 1, 1), (0.75, 0.9, 0.9, 0.95, 0.9, 0.9)),
    "VG": ((0.9, 1.0, 1.0, 1.0, 1, 1), (0.95, 1.0, 1.0, 0.95, 0.9, 0.9)),
}
COCOSO_FULL = {
    "VP": "Very Poor", "P": "Poor", "MP": "Medium Poor", "F": "Fair",
    "MG": "Medium Good", "G": "Good", "VG": "Very Good"
}

def _is_benefit_type(t):
    s = str(t).strip().lower()
    return (s.startswith("b") or "benefit" in s or "ben" == s or "max" in s)

def _safe_div(num, den):
    den = float(den)
    if den == 0:
        return 0.0
    return float(num) / den

def normalize_it2_matrix_excel(agg_matrix, criteria_types, alternatives, criteria):
    norm = {}
    for j, crit in enumerate(criteria):
        all_a, all_b, all_c, all_d = [], [], [], []
        for alt in alternatives:
            Au, _ = agg_matrix[(alt, crit)]
            all_a.append(float(Au[0]))
            all_b.append(float(Au[1]))
            all_c.append(float(Au[2]))
            all_d.append(float(Au[3]))
        delta_plus = max(max(all_a), max(all_b), max(all_c), max(all_d)) if alternatives else 1.0
        delta_minus = min(min(all_a), min(all_b), min(all_c), min(all_d)) if alternatives else 0.0
        if _is_benefit_type(criteria_types[j]):
            div = delta_plus if delta_plus != 0 else 1.0
            for alt in alternatives:
                Au, Al = agg_matrix[(alt, crit)]
                a, b, c, d, uh1, uh2 = Au
                e, f, g, h, lh1, lh2 = Al
                norm[(alt, crit)] = (
                    (_safe_div(a, div), _safe_div(b, div), _safe_div(c, div), _safe_div(d, div), uh1, uh2),
                    (_safe_div(e, div), _safe_div(f, div), _safe_div(g, div), _safe_div(h, div), lh1, lh2),
                )
        else:
            m = delta_minus
            for alt in alternatives:
                Au, Al = agg_matrix[(alt, crit)]
                aU, bU, cU, dU, uh1, uh2 = Au
                eL, fL, gL, hL, lh1, lh2 = Al
                norm_umf = (_safe_div(m, dU), _safe_div(m, cU), _safe_div(m, bU), _safe_div(m, aU), uh1, uh2)
                norm_lmf = (_safe_div(m, hL), _safe_div(m, gL), _safe_div(m, fL), _safe_div(m, eL), lh1, lh2)
                norm[(alt, crit)] = (norm_umf, norm_lmf)
    return norm

def it2_to_row(it2):
    Au, Al = it2
    return {
        "a": Au[0], "b": Au[1], "c": Au[2], "d": Au[3], "uh1": Au[4], "uh2": Au[5],
        "e": Al[0], "f": Al[1], "g": Al[2], "h": Al[3], "lh1": Al[4], "lh2": Al[5],
    }

def format_it2_table(matrix_dict, alternatives, criteria, value_formatter=format_it2):
    df = pd.DataFrame(index=alternatives, columns=criteria, dtype=object)
    for alt in alternatives:
        for crit in criteria:
            df.loc[alt, crit] = value_formatter(matrix_dict[(alt, crit)])
    return df

# ============================================================
# Delphi functions
# ============================================================
DELPHI_NUMERIC_SCALE = {
    1: ((0, 0.1, 0.1, 0.1, 1, 1), (0.0, 0.1, 0.1, 0.05, 0.9, 0.9)),
    2: ((0.2, 0.3, 0.3, 0.4, 1, 1), (0.25, 0.3, 0.3, 0.35, 0.9, 0.9)),
    3: ((0.4, 0.5, 0.5, 0.6, 1, 1), (0.45, 0.5, 0.5, 0.55, 0.9, 0.9)),
    4: ((0.6, 0.7, 0.7, 0.8, 1, 1), (0.65, 0.7, 0.7, 0.75, 0.9, 0.9)),
    5: ((0.8, 0.9, 0.9, 1.0, 1, 1), (0.85, 0.9, 0.9, 0.95, 0.9, 0.9)),
}
DELPHI_TERM_LABELS = {1: "VLR", 2: "LR", 3: "MR", 4: "HR", 5: "VHR"}
DELPHI_TERM_FULL = {1: "Very Low Relevance", 2: "Low Relevance", 3: "Medium Relevance", 4: "High Relevance", 5: "Very High Relevance"}

def parse_loose_numeric(x):
    if pd.isna(x):
        return np.nan
    if isinstance(x, (int, float, np.integer, np.floating)):
        return float(x)
    s = str(x).strip()
    m = re.search(r"-?\d+(?:\.\d+)?", s)
    if m:
        try:
            return float(m.group())
        except Exception:
            return np.nan
    return np.nan

def build_weight_vector(weight_series):
    w = pd.Series(weight_series).apply(parse_loose_numeric).fillna(0.0).astype(float)
    w = w.clip(lower=0.0)
    if float(w.sum()) <= 0:
        return np.ones(len(w)) / len(w) if len(w) > 0 else np.array([])
    return (w / w.sum()).to_numpy(dtype=float)

def aggregate_it2_delphi_from_scores(scores_df):
    aggregated = {}
    summary_rows = []
    for crit in scores_df.columns:
        col = scores_df[crit].apply(parse_loose_numeric)
        valid_scores = col[col.isin([1, 2, 3, 4, 5])].astype(int).tolist()
        if len(valid_scores) == 0:
            aggregated[crit] = zero_it2()
            summary_rows.append({"Criterion": crit, "N valid": 0, "Mean rating": np.nan, "Std rating": np.nan, "Mode rating": np.nan,
                "a": np.nan, "b": np.nan, "c": np.nan, "d": np.nan, "uh1": np.nan, "uh2": np.nan,
                "e": np.nan, "f": np.nan, "g": np.nan, "h": np.nan, "lh1": np.nan, "lh2": np.nan,
                "Score(UMF)": np.nan, "Score(LMF)": np.nan, "Crisp score": np.nan})
            continue
        n_valid = len(valid_scores)
        umf_a = sum(DELPHI_NUMERIC_SCALE[s][0][0] for s in valid_scores) / n_valid
        umf_b = sum(DELPHI_NUMERIC_SCALE[s][0][1] for s in valid_scores) / n_valid
        umf_c = sum(DELPHI_NUMERIC_SCALE[s][0][2] for s in valid_scores) / n_valid
        umf_d = sum(DELPHI_NUMERIC_SCALE[s][0][3] for s in valid_scores) / n_valid
        lmf_e = sum(DELPHI_NUMERIC_SCALE[s][1][0] for s in valid_scores) / n_valid
        lmf_f = sum(DELPHI_NUMERIC_SCALE[s][1][1] for s in valid_scores) / n_valid
        lmf_g = sum(DELPHI_NUMERIC_SCALE[s][1][2] for s in valid_scores) / n_valid
        lmf_h = sum(DELPHI_NUMERIC_SCALE[s][1][3] for s in valid_scores) / n_valid
        uh1 = min(DELPHI_NUMERIC_SCALE[s][0][4] for s in valid_scores)
        uh2 = min(DELPHI_NUMERIC_SCALE[s][0][5] for s in valid_scores)
        lh1 = min(DELPHI_NUMERIC_SCALE[s][1][4] for s in valid_scores)
        lh2 = min(DELPHI_NUMERIC_SCALE[s][1][5] for s in valid_scores)
        agg_it2 = ((umf_a, umf_b, umf_c, umf_d, uh1, uh2), (lmf_e, lmf_f, lmf_g, lmf_h, lh1, lh2))
        aggregated[crit] = agg_it2
        score_u, score_l, crisp = it2_score_components(agg_it2)
        valid_numeric = np.array(valid_scores, dtype=float)
        mode_val = pd.Series(valid_numeric).mode()
        mode_rating = float(mode_val.iloc[0]) if not mode_val.empty else np.nan
        summary_rows.append({"Criterion": crit, "N valid": len(valid_scores), "Mean rating": float(np.mean(valid_numeric)),
            "Std rating": float(np.std(valid_numeric, ddof=1)) if len(valid_numeric) > 1 else 0.0, "Mode rating": mode_rating,
            "a": umf_a, "b": umf_b, "c": umf_c, "d": umf_d, "uh1": uh1, "uh2": uh2,
            "e": lmf_e, "f": lmf_f, "g": lmf_g, "h": lmf_h, "lh1": lh1, "lh2": lh2,
            "Score(UMF)": score_u, "Score(LMF)": score_l, "Crisp score": crisp})
    summary_df = pd.DataFrame(summary_rows)
    return aggregated, summary_df

def detect_delphi_columns(df):
    candidates = []
    for col in df.columns:
        col_name = str(col)
        if col_name.lower().startswith("unnamed"):
            continue
        s = pd.to_numeric(df[col], errors="coerce")
        first_chunk = s.iloc[:100].dropna()
        if len(first_chunk) == 0:
            continue
        ratio_1_to_5 = first_chunk.between(1, 5, inclusive="both").mean()
        if ratio_1_to_5 >= 0.80:
            candidates.append(col)
    return candidates

def extract_contiguous_response_block(df, criteria_cols):
    if len(criteria_cols) == 0:
        return df.iloc[0:0].copy()
    temp = df[criteria_cols].apply(pd.to_numeric, errors="coerce")
    row_has_response = temp.notna().any(axis=1)
    started = False
    keep_idx = []
    for idx, flag in row_has_response.items():
        if flag and not started:
            started = True
        if started:
            if flag:
                keep_idx.append(idx)
            else:
                break
    if len(keep_idx) == 0:
        return df.iloc[0:0].copy()
    return df.loc[keep_idx].copy()

def format_delphi_scale_table():
    rows = []
    for k, v in DELPHI_NUMERIC_SCALE.items():
        rows.append({"Score": k, "Abbreviation": DELPHI_TERM_LABELS[k], "Meaning": DELPHI_TERM_FULL[k], "IT2TrFS": format_it2(v)})
    return pd.DataFrame(rows)

def render_delphi_results(clean_scores_df, summary_df, threshold):
    st.markdown("#### Cleaned response matrix")
    clean_it2_df = clean_scores_df.copy().map(lambda x: format_it2(DELPHI_NUMERIC_SCALE[int(x)]) if pd.notna(x) and int(x) in DELPHI_NUMERIC_SCALE else "")
    st.dataframe(clean_it2_df, use_container_width=True)
    out = summary_df.copy()
    out["Aggregated IT2TrFS"] = out.apply(lambda row: format_it2(((row["a"], row["b"], row["c"], row["d"], row["uh1"], row["uh2"]), (row["e"], row["f"], row["g"], row["h"], row["lh1"], row["lh2"]))), axis=1)
    out["Status"] = np.where(out["Crisp score"] >= threshold, "Accepted", "Rejected")
    out["Rank"] = out["Crisp score"].rank(ascending=False, method="min").astype("Int64")
    out = out.sort_values(["Status", "Crisp score"], ascending=[False, False]).reset_index(drop=True)
    display_cols = ["Criterion", "Aggregated IT2TrFS", "Mode rating", "Score(UMF)", "Score(LMF)", "Crisp score", "Status", "Rank"]
    display_df = out[display_cols].copy()
    st.markdown("#### Aggregated IT2TrFS-Delphi results")
    st.dataframe(display_df.style.format({"Mode rating": "{:.0f}", "Score(UMF)": "{:.6f}", "Score(LMF)": "{:.6f}", "Crisp score": "{:.6f}"}), use_container_width=True, hide_index=True)
    accepted = out.loc[out["Status"] == "Accepted", "Criterion"].tolist()
    rejected = out.loc[out["Status"] == "Rejected", "Criterion"].tolist()
    c1, c2, c3 = st.columns(3)
    c1.metric("Accepted criteria", len(accepted))
    c2.metric("Rejected criteria", len(rejected))
    c3.metric("Threshold", f"{threshold:.2f}")
    if accepted:
        st.success("Accepted criteria:\n\n" + "\n".join([f"- {x}" for x in accepted]))
    else:
        st.warning("No criterion passed the current threshold.")
    if rejected:
        with st.expander("Rejected criteria"):
            st.write("\n".join([f"- {x}" for x in rejected]))
    st.bar_chart(out.set_index("Criterion")["Crisp score"])
    csv_bytes = display_df.to_csv(index=False).encode("utf-8")
    st.download_button("⬇️ Download Delphi summary as CSV", data=csv_bytes, file_name="it2trfs_delphi_summary.csv", mime="text/csv", use_container_width=True)
    return out, accepted

def delphi_app():
    render_page_banner(
    "📘 IT2TrFS-Delphi Screening Module",
    "Standalone IT2TrFS-Delphi module. No expert-weight input is used in this page.",
    theme="default"
    )
    tab_howto, tab_excel, tab_manual = st.tabs(["📘 How to Use", "📂 Excel Upload", "✍️ Manual Entry"])
    with tab_howto:
        st.markdown("**What this page does**\n- Converts Delphi scores (1–5) into IT2TrFS values.\n- Aggregates responses criterion-wise using equal contribution from all valid respondents.\n- Computes Score(UMF), Score(LMF), and the final crisp score.\n- Screens criteria using a user-defined acceptance threshold.")
        st.markdown("**Numeric-to-IT2TrFS scale used in this Delphi page**")
        st.dataframe(format_delphi_scale_table(), hide_index=True, use_container_width=True)
        st.info("This Delphi page runs independently and does not apply expert weights.")
    with tab_excel:
        uploaded_file = st.file_uploader("Upload Delphi response Excel file", type=["xlsx", "xls"], key="delphi_excel_upload")
        threshold = st.number_input("Acceptance threshold for crisp score", min_value=0.0, max_value=1.0, value=0.60, step=0.01, key="delphi_threshold_excel")
        if uploaded_file is not None:
            try:
                xls = pd.ExcelFile(uploaded_file)
                sheet_name = st.selectbox("Sheet", options=xls.sheet_names, index=0, key="delphi_sheet_select")
                df_raw = pd.read_excel(uploaded_file, sheet_name=sheet_name)
                st.markdown("#### Preview of uploaded sheet")
                st.dataframe(df_raw.head(10), use_container_width=True)
                detected_cols = detect_delphi_columns(df_raw)
                response_block = extract_contiguous_response_block(df_raw, detected_cols)
                with st.expander("Advanced settings"):
                    criteria_cols = st.multiselect("Detected Delphi criteria columns", options=list(df_raw.columns), default=detected_cols, key="delphi_criteria_cols")
                    start_row = st.number_input("Start row of response block (1-based, after header)", min_value=1, max_value=max(len(df_raw), 1), value=1 if response_block.empty else int(response_block.index.min()) + 2, step=1, key="delphi_start_row")
                    end_row = st.number_input("End row of response block (1-based, after header)", min_value=1, max_value=max(len(df_raw), 1), value=min(len(df_raw), int(response_block.index.max()) + 2) if not response_block.empty else min(10, len(df_raw)), step=1, key="delphi_end_row")
                if len(criteria_cols) == 0:
                    st.warning("No Delphi criteria columns selected.")
                else:
                    response_block = df_raw.iloc[int(start_row) - 1:int(end_row)].copy()
                    scores_df = response_block[criteria_cols].map(parse_loose_numeric)
                    scores_df = scores_df.where(scores_df.isin([1, 2, 3, 4, 5]), np.nan)
                    valid_rows = scores_df.notna().any(axis=1)
                    scores_df = scores_df.loc[valid_rows].reset_index(drop=True)
                    st.caption(f"Detected/selected responses: {len(scores_df)} | Criteria: {len(criteria_cols)}")
                    if st.button("✅ Run IT2TrFS-Delphi from Excel", type="primary", use_container_width=True, key="delphi_excel_run"):
                        if len(scores_df) == 0:
                            st.error("No valid Delphi responses (1–5 scale) were found in the selected block.")
                        else:
                            _, summary_df = aggregate_it2_delphi_from_scores(scores_df)
                            render_delphi_results(scores_df, summary_df, threshold)
            except Exception as e:
                st.error(f"Could not read the workbook: {e}")
    with tab_manual:
        c1, c2 = st.columns(2)
        n_experts = c1.number_input("Number of respondents", min_value=1, max_value=50, value=5, step=1, key="delphi_manual_nexp")
        n_criteria = c2.number_input("Number of criteria", min_value=1, max_value=50, value=6, step=1, key="delphi_manual_ncrit")
        threshold_manual = st.number_input("Acceptance threshold", min_value=0.0, max_value=1.0, value=0.60, step=0.01, key="delphi_threshold_manual")
        default_criteria = [f"C{i+1}" for i in range(n_criteria)]
        criteria_names = []
        cols = st.columns(min(4, n_criteria))
        for i in range(n_criteria):
            with cols[i % len(cols)]:
                criteria_names.append(st.text_input(f"Criterion {i+1}", value=default_criteria[i], key=f"delphi_crit_{i}"))
        expected_index = [f"R{i+1}" for i in range(n_experts)]
        if "delphi_manual_df" not in st.session_state:
            st.session_state.delphi_manual_df = pd.DataFrame(3, index=expected_index, columns=criteria_names)
        old_df = st.session_state.delphi_manual_df
        if list(old_df.index) != expected_index or list(old_df.columns) != criteria_names:
            st.session_state.delphi_manual_df = pd.DataFrame(3, index=expected_index, columns=criteria_names)
        st.markdown("**Manual Delphi score matrix (1–5)**")
        edited_scores = st.data_editor(st.session_state.delphi_manual_df, use_container_width=True, column_config={c: st.column_config.NumberColumn(c, min_value=1, max_value=5, step=1, format="%d") for c in criteria_names}, key="delphi_manual_editor")
        if st.button("✅ Run manual IT2TrFS-Delphi", type="primary", use_container_width=True, key="delphi_manual_run"):
            clean_scores = edited_scores.map(parse_loose_numeric)
            clean_scores = clean_scores.where(clean_scores.isin([1, 2, 3, 4, 5]), np.nan)
            _, summary_df = aggregate_it2_delphi_from_scores(clean_scores)
            render_delphi_results(clean_scores, summary_df, threshold_manual)

# ============================================================
# CoCoSo App
# ============================================================
def cocoso_app():
    render_page_banner(
    "📊 IT2TrFS-CoCoSo",
    "Normalization uses δ⁺ (max) for Benefit and δ⁻ (min) for Cost.",
    theme="green"
    )
    with st.expander("Linguistic scale (VP…VG)"):
        scale_df = pd.DataFrame([{"Abbr": k, "Meaning": COCOSO_FULL[k], "IT2TrFS": format_it2(v)} for k, v in COCOSO_LINGUISTIC_TERMS.items()])
        st.dataframe(scale_df, hide_index=True, use_container_width=True)
    st.subheader("Step 1: Alternatives, Criteria, Types, Weights")
    c1, c2 = st.columns(2)
    alts_in = c1.text_input("Alternatives (comma-separated)", "T1, T2, T3", key="cocoso_alts_in")
    crits_in = c2.text_input("Criteria (comma-separated)", "C1, C2, C3", key="cocoso_crits_in")
    alternatives = [a.strip() for a in alts_in.split(",") if a.strip()]
    criteria = [c.strip() for c in crits_in.split(",") if c.strip()]
    if len(alternatives) < 1 or len(criteria) < 1:
        st.warning("Please provide at least 1 alternative and 1 criterion.")
        return
    if "cocoso_crit_df_it2" not in st.session_state or list(st.session_state.cocoso_crit_df_it2.get("Criterion", [])) != criteria:
        w = [round(1 / len(criteria), 6)] * len(criteria)
        if len(criteria) > 0:
            w[-1] = 1.0 - sum(w[:-1])
        st.session_state.cocoso_crit_df_it2 = pd.DataFrame({"Criterion": criteria, "Type": ["Benefit"] * len(criteria), "Weight": w})
    edited_crit_df = st.data_editor(st.session_state.cocoso_crit_df_it2, hide_index=True, use_container_width=True, column_config={"Type": st.column_config.SelectboxColumn("Type", options=["Benefit", "Cost"]), "Weight": st.column_config.NumberColumn("Weight", format="%.5f", min_value=0.0, max_value=1.0, step=0.00001)}, key="cocoso_crit_editor_it2")
    criteria_types = edited_crit_df["Type"].tolist()
    criteria_weights = edited_crit_df["Weight"].astype(float).tolist()
    if not np.isclose(sum(criteria_weights), 1.0):
        st.error(f"Criteria weights must sum to 1.0 (now: {sum(criteria_weights):.5f}).")
        return
    st.subheader("Step 2: Expert evaluations (linguistic)")
    num_experts = st.number_input("Number of experts", min_value=1, max_value=30, value=2, step=1, key="cocoso_ne_it2")
    st.markdown("**Expert weights** (must sum to 1.0)")
    expert_weights = []
    if num_experts == 1:
        expert_weights = [1.0]
        st.info("Single expert → weight = 1.0")
    else:
        cols = st.columns(num_experts)
        for i in range(num_experts):
            with cols[i]:
                expert_weights.append(st.number_input(f"E{i+1}", min_value=0.0, max_value=1.0, value=round(1 / num_experts, 6), step=0.01, format="%.6f", key=f"cocoso_ew_{i}"))
        if not np.isclose(sum(expert_weights), 1.0):
            st.error(f"Expert weights must sum to 1.0 (now: {sum(expert_weights):.5f}).")
            return
    if "cocoso_expert_dfs_it2" not in st.session_state:
        st.session_state.cocoso_expert_dfs_it2 = {}
    need_reset = (len(st.session_state.cocoso_expert_dfs_it2) != num_experts) or (num_experts > 0 and (list(st.session_state.cocoso_expert_dfs_it2.get(0, pd.DataFrame()).index) != alternatives or list(st.session_state.cocoso_expert_dfs_it2.get(0, pd.DataFrame()).columns) != criteria))
    if need_reset:
        st.session_state.cocoso_expert_dfs_it2 = {i: pd.DataFrame("F", index=alternatives, columns=criteria) for i in range(num_experts)}
    tabs = st.tabs([f"Expert {i+1}" for i in range(num_experts)])
    for i, tab in enumerate(tabs):
        with tab:
            st.session_state.cocoso_expert_dfs_it2[i] = st.data_editor(st.session_state.cocoso_expert_dfs_it2[i], use_container_width=True, column_config={c: st.column_config.SelectboxColumn(c, options=list(COCOSO_LINGUISTIC_TERMS.keys())) for c in criteria}, key=f"cocoso_editor_it2_{i}")
    st.subheader("Step 3: Calculate")
    tau = st.number_input("τ (tau)", min_value=0.0, max_value=1.0, value=0.5, step=0.05, key="cocoso_tau")
    if st.button("✅ Run IT2TrFS-CoCoSo", type="primary", use_container_width=True, key="cocoso_run_it2"):
        with st.spinner("Computing..."):
            agg_matrix = {}
            for alt in alternatives:
                for crit in criteria:
                    acc = None
                    for e in range(num_experts):
                        term = st.session_state.cocoso_expert_dfs_it2[e].loc[alt, crit]
                        it2 = COCOSO_LINGUISTIC_TERMS[term]
                        it2w = scalar_mul_it2(expert_weights[e], it2)
                        acc = it2w if acc is None else add_it2(acc, it2w)
                    agg_matrix[(alt, crit)] = acc
            st.markdown("#### 3.1 Aggregated IT2TrFS Decision Matrix")
            st.dataframe(format_it2_table(agg_matrix, alternatives, criteria), use_container_width=True)
            norm_matrix = normalize_it2_matrix_excel(agg_matrix=agg_matrix, criteria_types=criteria_types, alternatives=alternatives, criteria=criteria)
            st.markdown("#### 3.2 Normalized IT2TrFS Matrix")
            st.dataframe(format_it2_table(norm_matrix, alternatives, criteria), use_container_width=True)
            SBi = {}
            PBi = {}
            for alt in alternatives:
                s_acc = zero_it2()
                p_acc = zero_it2()
                for j, crit in enumerate(criteria):
                    r = norm_matrix[(alt, crit)]
                    wj = float(criteria_weights[j])
                    s_acc = add_it2(s_acc, scalar_mul_it2(wj, r))
                    p_acc = add_it2(p_acc, it2_pow(r, wj))
                SBi[alt] = s_acc
                PBi[alt] = p_acc
            sbi_df = pd.DataFrame([{"Alternative": alt, **it2_to_row(SBi[alt])} for alt in alternatives])
            pbi_df = pd.DataFrame([{"Alternative": alt, **it2_to_row(PBi[alt])} for alt in alternatives])
            st.markdown("#### 3.3 SBi (IT2TrFS)")
            st.dataframe(sbi_df.style.format(precision=6), use_container_width=True, hide_index=True)
            st.markdown("#### 3.3 PBi (IT2TrFS)")
            st.dataframe(pbi_df.style.format(precision=6), use_container_width=True, hide_index=True)
            crisp_S = {alt: cocoso_crisp_score(SBi[alt]) for alt in alternatives}
            crisp_P = {alt: cocoso_crisp_score(PBi[alt]) for alt in alternatives}
            df_crisp = pd.DataFrame({"Alternative": alternatives, "Crisp SBi": [crisp_S[a] for a in alternatives], "Crisp PBi": [crisp_P[a] for a in alternatives]})
            st.markdown("#### 3.4 Crisp SBi & Crisp PBi")
            st.dataframe(df_crisp.style.format(precision=6), use_container_width=True, hide_index=True)
            sumS = sum(crisp_S.values()); sumP = sum(crisp_P.values()); minS = min(crisp_S.values()); minP = min(crisp_P.values()); maxS = max(crisp_S.values()); maxP = max(crisp_P.values())
            rows = []
            denom_kic = (tau * maxS + (1.0 - tau) * maxP)
            denom_kic = denom_kic if denom_kic != 0 else 1.0
            for alt in alternatives:
                S = crisp_S[alt]; P = crisp_P[alt]
                Kia = (S + P) / (sumS + sumP) if (sumS + sumP) != 0 else 0.0
                Kib = (S / minS if minS != 0 else 0.0) + (P / minP if minP != 0 else 0.0)
                Kic = ((tau * S) + ((1.0 - tau) * P)) / denom_kic
                K = (Kia * Kib * Kic) ** (1 / 3) + ((Kia + Kib + Kic) / 3)
                rows.append({"Alternative": alt, "Kia": Kia, "Kib": Kib, "Kic": Kic, "K": K})
            dfK = pd.DataFrame(rows)
            dfK["Rank"] = dfK["K"].rank(ascending=False, method="min").astype(int)
            dfK = dfK.sort_values("Rank").reset_index(drop=True)
            st.markdown("#### 3.5 Final CoCoSo indices & Rank")
            st.dataframe(dfK.style.format(precision=6), use_container_width=True, hide_index=True)
                        # -------------------------
            # Excel export for CoCoSo
            # -------------------------
            agg_export = format_it2_table(agg_matrix, alternatives, criteria).reset_index().rename(columns={"index": "Alternative"})
            norm_export = format_it2_table(norm_matrix, alternatives, criteria).reset_index().rename(columns={"index": "Alternative"})

            crisp_export = df_crisp.copy()
            final_export = dfK.copy()

            excel_bytes = dataframe_dict_to_excel_bytes({
                "Aggregated Matrix": agg_export,
                "Normalized Matrix": norm_export,
                "SBi": sbi_df,
                "PBi": pbi_df,
                "Crisp Scores": crisp_export,
                "Final Ranking": final_export,
            })

            st.download_button(
                "⬇️ Download CoCoSo Results (Excel)",
                data=excel_bytes,
                file_name="it2trfs_cocoso_results.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                key="cocoso_excel_download"
            )
# ============================================================
# WINGS functions
# ============================================================
LINGUISTIC_TERMS = {
    "strength": {"VLR": ((0, 0.1, 0.1, 0.1, 1, 1), (0.0, 0.1, 0.1, 0.05, 0.9, 0.9)),
                 "LR": ((0.2, 0.3, 0.3, 0.4, 1, 1), (0.25, 0.3, 0.3, 0.35, 0.9, 0.9)),
                 "MR": ((0.4, 0.5, 0.5, 0.6, 1, 1), (0.45, 0.5, 0.5, 0.55, 0.9, 0.9)),
                 "HR": ((0.6, 0.7, 0.7, 0.8, 1, 1), (0.65, 0.7, 0.7, 0.75, 0.9, 0.9)),
                 "VHR": ((0.8, 0.9, 0.9, 1.0, 1, 1), (0.85, 0.90, 0.90, 0.95, 0.9, 0.9))},
    "influence": {"ELI": ((0, 0.1, 0.1, 0.2, 1, 1), (0.05, 0.1, 0.1, 0.15, 0.9, 0.9)),
                  "VLI": ((0.1, 0.2, 0.2, 0.35, 1, 1), (0.15, 0.2, 0.2, 0.3, 0.9, 0.9)),
                  "LI": ((0.2, 0.35, 0.35, 0.5, 1, 1), (0.25, 0.35, 0.35, 0.45, 0.9, 0.9)),
                  "MI": ((0.35, 0.5, 0.5, 0.65, 1, 1), (0.40, 0.5, 0.5, 0.6, 0.9, 0.9)),
                  "HI": ((0.5, 0.65, 0.65, 0.8, 1, 1), (0.55, 0.65, 0.65, 0.75, 0.9, 0.9)),
                  "VHI": ((0.65, 0.80, 0.80, 0.9, 1, 1), (0.7, 0.8, 0.8, 0.85, 0.9, 0.9)),
                  "EHI": ((0.8, 0.9, 0.9, 1.0, 1, 1), (0.85, 0.9, 0.9, 0.95, 0.9, 0.9))}
}
FULL_FORMS = {"VLR": "Very Low Relevance", "LR": "Low Relevance", "MR": "Medium Relevance", "HR": "High Relevance", "VHR": "Very High Relevance", "ELI": "Extremely Low Influence", "VLI": "Very Low Influence", "LI": "Low Influence", "MI": "Medium Influence", "HI": "High Influence", "VHI": "Very High Influence", "EHI": "Extremely High Influence"}

def defuzz_it2(A):
    Au, Al = A
    a, b, c, d, uh1, uh2 = Au
    e, f, g, h, lh1, lh2 = Al
    score_u = (((d - a) + ((uh2 * c) - a) + ((uh1 * b) - a)) / 4.0) + a
    score_l = (((h - e) + ((lh2 * g) - e) + ((lh1 * f) - e)) / 4.0) + e
    return (score_u + score_l) / 2.0

def identity_it2(n):
    I_mat = [[zero_it2() for _ in range(n)] for _ in range(n)]
    for i in range(n):
        I_mat[i][i] = ((1, 1, 1, 1, 1, 1), (1, 1, 1, 1, 1, 1))
    return I_mat

def compute_total_relation_matrix(normalized_matrix):
    n = len(normalized_matrix)
    Z_4d = np.zeros((2, 2, n, n, 4))
    for i in range(n):
        for j in range(n):
            Au, Al = normalized_matrix[i][j]
            Z_4d[0, 0, i, j, :] = Au[:4]
            Z_4d[0, 1, i, j, :2] = Au[4:]
            Z_4d[1, 0, i, j, :] = Al[:4]
            Z_4d[1, 1, i, j, :2] = Al[4:]
    for i in range(2):
        for j in range(2):
            if j == 0:
                for k in range(4):
                    Z_component = Z_4d[i, j, :, :, k]
                    try:
                        T_component = Z_component @ np.linalg.pinv(np.eye(n) - Z_component)
                    except np.linalg.LinAlgError:
                        T_component = np.zeros((n, n))
                    Z_4d[i, j, :, :, k] = T_component
    T = [[zero_it2() for _ in range(n)] for _ in range(n)]
    for i in range(n):
        for j in range(n):
            T[i][j] = ((Z_4d[0, 0, i, j, 0], Z_4d[0, 0, i, j, 1], Z_4d[0, 0, i, j, 2], Z_4d[0, 0, i, j, 3], Z_4d[0, 1, i, j, 0], Z_4d[0, 1, i, j, 1]), (Z_4d[1, 0, i, j, 0], Z_4d[1, 0, i, j, 1], Z_4d[1, 0, i, j, 2], Z_4d[1, 0, i, j, 3], Z_4d[1, 1, i, j, 0], Z_4d[1, 1, i, j, 1]))
    return T

def calculate_TI_TR(T):
    n = len(T)
    TI = [zero_it2() for _ in range(n)]
    TR = [zero_it2() for _ in range(n)]
    for i in range(n):
        for j in range(n):
            TI[i] = add_it2(TI[i], T[i][j])
            TR[j] = add_it2(TR[j], T[i][j])
    return TI, TR

def wings_method_experts(strengths_list, influence_matrices_list, weights=None):
    n = len(strengths_list[0])
    num_experts = len(strengths_list)
    if weights is None:
        weights = [1.0 / num_experts] * num_experts
    avg_sidrm = [[zero_it2() for _ in range(n)] for _ in range(n)]
    for exp in range(num_experts):
        w = weights[exp]
        for i in range(n):
            str_w = scalar_mul_it2(w, strengths_list[exp][i])
            avg_sidrm[i][i] = add_it2(avg_sidrm[i][i], str_w)
            for j in range(n):
                if i != j:
                    inf_w = scalar_mul_it2(w, influence_matrices_list[exp][i][j])
                    avg_sidrm[i][j] = add_it2(avg_sidrm[i][j], inf_w)
    s1U = s2U = s3U = s4U = s1L = s2L = s3L = s4L = 0.0
    for i in range(n):
        for j in range(n):
            Au, Al = avg_sidrm[i][j]
            s1U += Au[0]; s2U += Au[1]; s3U += Au[2]; s4U += Au[3]
            s1L += Al[0]; s2L += Al[1]; s3L += Al[2]; s4L += Al[3]
    s = s1U + s2U + s3U + s4U + s1L + s2L + s3L + s4L
    Z_mat = [[zero_it2() for _ in range(n)] for _ in range(n)]
    for i in range(n):
        for j in range(n):
            Au, Al = avg_sidrm[i][j]
            new_u = (Au[0]/s if s else 0.0, Au[1]/s if s else 0.0, Au[2]/s if s else 0.0, Au[3]/s if s else 0.0, Au[4], Au[5])
            new_l = (Al[0]/s if s else 0.0, Al[1]/s if s else 0.0, Al[2]/s if s else 0.0, Al[3]/s if s else 0.0, Al[4], Al[5])
            Z_mat[i][j] = (new_u, new_l)
    T_mat = compute_total_relation_matrix(Z_mat)
    TI, TR = calculate_TI_TR(T_mat)
    engagement = [add_it2(TI[i], TR[i]) for i in range(n)]
    role = [sub_it2(TI[i], TR[i]) for i in range(n)]
    TI_defuzz = np.array([defuzz_it2(TI[i]) for i in range(n)], dtype=float)
    TR_defuzz = np.array([defuzz_it2(TR[i]) for i in range(n)], dtype=float)
    engagement_defuzz = np.array([defuzz_it2(engagement[i]) for i in range(n)], dtype=float)
    role_defuzz = np.array([defuzz_it2(role[i]) for i in range(n)], dtype=float)
    Ew = np.sqrt(np.square(engagement_defuzz) + np.square(role_defuzz))
    sum_Ew = float(np.sum(Ew))
    normalized_weights = (Ew / sum_Ew) if sum_Ew > 0 else np.zeros_like(Ew)
    return {"average_sidrm": avg_sidrm, "scaling_factor": s, "normalized_matrix": Z_mat, "total_matrix": T_mat,
            "total_impact": TI, "total_receptivity": TR, "engagement": engagement, "role": role,
            "total_impact_defuzz": TI_defuzz, "total_receptivity_defuzz": TR_defuzz,
            "engagement_defuzz": engagement_defuzz, "role_defuzz": role_defuzz, "Ew": Ew, "normalized_weights": normalized_weights}

def format_it2_df(mat, index, columns):
    df = pd.DataFrame(index=index, columns=columns)
    for i in range(len(index)):
        for j in range(len(columns)):
            df.iloc[i, j] = format_it2(mat[i][j])
    return df

def generate_flowchart_for_expert(expert_data, component_names, expert_idx=None):
    graph = graphviz.Digraph(comment=f"IT2TrFS WINGS Flowchart - Expert {expert_idx+1}" if expert_idx is not None else "IT2TrFS WINGS Flowchart")
    graph.attr(rankdir="TD", size="8,8")
    for comp_idx, comp_name in enumerate(component_names):
        strength = expert_data["strengths_linguistic"][comp_idx]
        label = f"{comp_name} ({strength})"
        graph.node(comp_name, label=label, shape="box", style="rounded,filled", fillcolor="lightblue", fontsize="12")
    for from_idx, from_comp in enumerate(component_names):
        for to_idx, to_comp in enumerate(component_names):
            if from_idx == to_idx:
                continue
            influence = expert_data["influence_matrix_linguistic"][from_idx][to_idx]
            if influence != "ELI":
                graph.edge(from_comp, to_comp, label=influence)
    return graph

def create_word_report(results, component_names, n_experts=1, expert_weights=None):
    doc = Document()
    title = doc.add_heading("IT2TrFS WINGS Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from datetime import datetime
    doc.add_paragraph(f"Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Number of experts: {n_experts}")
    if expert_weights and n_experts > 1:
        weights_text = "Expert weights: " + ", ".join([f"Expert {i+1}: {weight:.4f}" for i, weight in enumerate(expert_weights)])
        doc.add_paragraph(weights_text)
    comp_para = doc.add_paragraph("Components analyzed: ")
    for i, name in enumerate(component_names):
        comp_para.add_run(f"{i+1}. {name}  ")
    doc.add_heading("Impact, Receptivity, Engagement, Role, and Normalized Weight Results", level=1)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Component"; hdr[1].text = "Total Impact (TI)"; hdr[2].text = "Total Receptivity (TR)"; hdr[3].text = "TI + TR"; hdr[4].text = "TI - TR"; hdr[5].text = "E(wj)"; hdr[6].text = "Normalized Weight"
    for i, name in enumerate(component_names):
        row = table.add_row().cells
        row[0].text = name
        row[1].text = f"{results['total_impact_defuzz'][i]:.6f}"
        row[2].text = f"{results['total_receptivity_defuzz'][i]:.6f}"
        row[3].text = f"{results['engagement_defuzz'][i]:.6f}"
        row[4].text = f"{results['role_defuzz'][i]:.6f}"
        row[5].text = f"{results['Ew'][i]:.6f}"
        row[6].text = f"{results['normalized_weights'][i]:.6f}"
    return doc

def get_word_download_link(doc):
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    b64 = base64.b64encode(file_stream.read()).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="it2trfs_wings_analysis_report.docx">Download Word Report</a>'
    return href

def wings_app():
    render_page_banner(
        "📊 IT2TrFS WINGS Method Analysis Platform",
        "IT2TrFS-WINGS module for causal analysis, interaction mapping, and weight derivation.",
        theme="orange"
    )

    tab_howto, tab_analysis = st.tabs(["📘 How to Use", "📊 Analysis"])

    with tab_howto:
        st.markdown("Use the sidebar to configure components/experts and run IT2TrFS-WINGS.")
        with st.expander("Linguistic Terms Reference"):
            col1, col2 = st.columns(2)

            with col1:
                st.write("**Strength/Relevance Terms**")
                strength_df = pd.DataFrame([
                    {
                        "Abbreviation": abbr,
                        "Full Form": FULL_FORMS[abbr],
                        "IT2TrFS": format_it2(it2)
                    }
                    for abbr, it2 in LINGUISTIC_TERMS["strength"].items()
                ])
                st.dataframe(strength_df, hide_index=True, use_container_width=True)

            with col2:
                st.write("**Influence Terms**")
                infl_df = pd.DataFrame([
                    {
                        "Abbreviation": abbr,
                        "Full Form": FULL_FORMS[abbr],
                        "IT2TrFS": format_it2(it2)
                    }
                    for abbr, it2 in LINGUISTIC_TERMS["influence"].items()
                ])
                st.dataframe(infl_df, hide_index=True, use_container_width=True)

    with tab_analysis:
        with st.sidebar:
            st.header("⚙️ IT2TrFS-WINGS Configuration")
            n_components = st.number_input("Number of Components", min_value=2, max_value=25, value=3, key="w_ncomp")
            n_experts = st.number_input("Number of Experts", min_value=1, max_value=15, value=1, key="w_nexp")
            component_names = []
            for i in range(n_components):
                component_names.append(st.text_input(f"Name of Component {i+1}", value=f"C{i+1}", key=f"w_comp_{i}"))
            expert_weights = None
            if n_experts > 1:
                st.markdown("---")
                st.subheader("Expert Weights (sum=1)")
                weights = []
                for i in range(n_experts):
                    weights.append(st.number_input(f"Weight E{i+1}", min_value=0.0, max_value=1.0, value=round(1 / n_experts, 4), step=0.01, key=f"w_w_{i}"))
                if not np.isclose(sum(weights), 1.0):
                    st.error(f"Weights must sum to 1.0 (now: {sum(weights):.4f})")
                    st.stop()
                expert_weights = weights
        if "experts_data" not in st.session_state or not isinstance(st.session_state.experts_data, dict):
            st.session_state.experts_data = {}
        stale_expert_keys = [k for k in st.session_state.experts_data.keys() if isinstance(k, int) and k >= n_experts]
        for k in stale_expert_keys:
            del st.session_state.experts_data[k]
        for e in range(n_experts):
            existing = st.session_state.experts_data.get(e, {})
            old_strengths = existing.get("strengths_linguistic", [])
            if not isinstance(old_strengths, list):
                old_strengths = []
            refreshed_strengths = []
            for i in range(n_components):
                val = old_strengths[i] if i < len(old_strengths) else "HR"
                if val not in LINGUISTIC_TERMS["strength"]:
                    val = "HR"
                refreshed_strengths.append(val)
            old_infl = existing.get("influence_matrix_linguistic", [])
            if not isinstance(old_infl, list):
                old_infl = []
            refreshed_infl = []
            for i in range(n_components):
                row = old_infl[i] if i < len(old_infl) and isinstance(old_infl[i], list) else []
                new_row = []
                for j in range(n_components):
                    if i == j:
                        new_row.append("ELI")
                    else:
                        val = row[j] if j < len(row) else "ELI"
                        if val not in LINGUISTIC_TERMS["influence"]:
                            val = "ELI"
                        new_row.append(val)
                refreshed_infl.append(new_row)
            st.session_state.experts_data[e] = {"strengths_linguistic": refreshed_strengths, "influence_matrix_linguistic": refreshed_infl}
        tabs = st.tabs([f"Expert {i+1}" for i in range(n_experts)]) if n_experts > 1 else [st.container()]
        strengths_list = []
        influence_list = []
        for e in range(n_experts):
            with tabs[e]:
                if n_experts > 1:
                    st.markdown(f"### Expert {e+1}")
                strengths = []
                st.write("**Strength / Relevance of Components**")
                cols = st.columns(n_components)
                for i in range(n_components):
                    with cols[i]:
                        cur = st.session_state.experts_data[e]["strengths_linguistic"][i] if i < len(st.session_state.experts_data[e]["strengths_linguistic"]) else "HR"
                        term = st.selectbox(component_names[i], options=list(LINGUISTIC_TERMS["strength"].keys()), index=list(LINGUISTIC_TERMS["strength"].keys()).index(cur), key=f"w_str_{e}_{i}")
                        st.session_state.experts_data[e]["strengths_linguistic"][i] = term
                        strengths.append(LINGUISTIC_TERMS["strength"][term])
                st.write("**Influence Matrix** (row influences column)")
                inf_mat = [[None] * n_components for _ in range(n_components)]
                header_cols = st.columns(n_components + 1)
                with header_cols[0]:
                    st.markdown("**From \\ To**")
                for j in range(n_components):
                    with header_cols[j + 1]:
                        st.markdown(f"**{component_names[j]}**")
                for i in range(n_components):
                    row_cols = st.columns(n_components + 1)
                    with row_cols[0]:
                        st.markdown(f"**{component_names[i]}**")
                    for j in range(n_components):
                        with row_cols[j + 1]:
                            if i == j:
                                st.markdown("—")
                                inf_mat[i][j] = zero_it2()
                                st.session_state.experts_data[e]["influence_matrix_linguistic"][i][j] = "ELI"
                            else:
                                cur = st.session_state.experts_data[e]["influence_matrix_linguistic"][i][j] if i < len(st.session_state.experts_data[e]["influence_matrix_linguistic"]) and j < len(st.session_state.experts_data[e]["influence_matrix_linguistic"][i]) else "ELI"
                                term = st.selectbox(f"{component_names[i]}→{component_names[j]}", options=list(LINGUISTIC_TERMS["influence"].keys()), index=list(LINGUISTIC_TERMS["influence"].keys()).index(cur), key=f"w_inf_{e}_{i}_{j}", label_visibility="collapsed")
                                st.session_state.experts_data[e]["influence_matrix_linguistic"][i][j] = term
                                inf_mat[i][j] = LINGUISTIC_TERMS["influence"][term]
                flowchart = generate_flowchart_for_expert(st.session_state.experts_data[e], component_names, expert_idx=e)
                st.graphviz_chart(flowchart, use_container_width=True)
                strengths_list.append(strengths)
                influence_list.append(inf_mat)
        st.markdown("---")
        if st.button("✅ Run IT2TrFS-WINGS", type="primary", use_container_width=True):
            results = wings_method_experts(strengths_list, influence_list, weights=expert_weights)
            t1, t2, t3 = st.tabs(["📋 Results", "📈 Plot", "📝 Export"])
            with t1:
                st.markdown("### Average SIDRM")
                st.dataframe(format_it2_df(results["average_sidrm"], component_names, component_names), use_container_width=True)
                st.markdown(f"**Scaling factor (s):** {results['scaling_factor']:.6f}")
                st.markdown("### Normalized Direct-Relation Matrix")
                st.dataframe(format_it2_df(results["normalized_matrix"], component_names, component_names), use_container_width=True)
                st.markdown("### Total Relation Matrix")
                st.dataframe(format_it2_df(results["total_matrix"], component_names, component_names), use_container_width=True)
                results_df = pd.DataFrame({"Component": component_names, "TI (IT2TrFS)": [format_it2(results["total_impact"][i]) for i in range(len(component_names))], "TR (IT2TrFS)": [format_it2(results["total_receptivity"][i]) for i in range(len(component_names))], "TI+TR (IT2TrFS)": [format_it2(results["engagement"][i]) for i in range(len(component_names))], "TI-TR (IT2TrFS)": [format_it2(results["role"][i]) for i in range(len(component_names))], "TI (Defuzzified)": results["total_impact_defuzz"], "TR (Defuzzified)": results["total_receptivity_defuzz"], "TI+TR": results["engagement_defuzz"], "TI-TR": results["role_defuzz"], "E(wj)": results["Ew"], "Normalized Weight": results["normalized_weights"]}).sort_values(by="Normalized Weight", ascending=False)
                st.markdown("### Final Results Table")
                st.dataframe(results_df.style.format({"TI (Defuzzified)": "{:.6f}", "TR (Defuzzified)": "{:.6f}", "TI+TR": "{:.6f}", "TI-TR": "{:.6f}", "E(wj)": "{:.6f}", "Normalized Weight": "{:.6f}"}), use_container_width=True, hide_index=True)
                classification_df = pd.DataFrame({"Component": component_names, "Type": ["Cause" if results["role_defuzz"][i] > 0 else "Effect" for i in range(len(component_names))], "TI-TR": results["role_defuzz"], "TI+TR": results["engagement_defuzz"], "E(wj)": results["Ew"], "Normalized Weight": results["normalized_weights"]}).sort_values(by="Normalized Weight", ascending=False)
                st.markdown("### Component Classification")
                st.dataframe(classification_df.style.format({"TI-TR": "{:.6f}", "TI+TR": "{:.6f}", "E(wj)": "{:.6f}", "Normalized Weight": "{:.6f}"}), use_container_width=True, hide_index=True)
            with t2:
                plot_df = pd.DataFrame({"Component": component_names, "TI+TR": results["engagement_defuzz"], "TI-TR": results["role_defuzz"], "Normalized Weight": results["normalized_weights"], "Type": ["Cause" if x > 0 else "Effect" for x in results["role_defuzz"]]})
                fig, ax = plt.subplots(figsize=(8, 6))
                for _, row in plot_df.iterrows():
                    ax.scatter(row["TI+TR"], row["TI-TR"], s=120)
                    ax.text(row["TI+TR"], row["TI-TR"], f" {row['Component']}", fontsize=10)
                ax.axhline(0, linestyle="--", linewidth=1)
                ax.axvline(plot_df["TI+TR"].mean(), linestyle="--", linewidth=1)
                ax.set_xlabel("TI + TR")
                ax.set_ylabel("TI - TR")
                ax.set_title("Cause-Effect Map")
                st.pyplot(fig)
                st.markdown("### Normalized Weight Summary")
                weight_df = pd.DataFrame({"Component": component_names, "Normalized Weight": results["normalized_weights"]}).sort_values(by="Normalized Weight", ascending=False)
                st.dataframe(weight_df.style.format({"Normalized Weight": "{:.6f}"}), use_container_width=True, hide_index=True)
            with t3:
                doc = create_word_report(results, component_names, n_experts, expert_weights)
                st.markdown(get_word_download_link(doc), unsafe_allow_html=True)


# ============================================================
# MAIN NAVIGATION
# ============================================================
def main():
    page = st.sidebar.radio(
        "Navigate",
        ["IT2TrFS-Delphi", "IT2TrFS-WINGS", "IT2TrFS-CoCoSo"],
    )

    st.sidebar.success("✅ Authenticated")

    if st.sidebar.button("🚪 Logout", use_container_width=True, type="secondary"):
        logout()
        st.rerun()

    # Main-body suite banner (right side)
    render_workspace_banner()

    if page == "IT2TrFS-Delphi":
        delphi_app()
    elif page == "IT2TrFS-WINGS":
        wings_app()
    elif page == "IT2TrFS-CoCoSo":
        cocoso_app()

    render_sidebar_research_profiles()
    render_footer()

# ============================================================
# ENTRY POINT
# ============================================================
if check_password():
    main()
else:
    st.stop()
